from PyQt5.QtWidgets import QMainWindow, QApplication
from tkinter.filedialog import askopenfilename, asksaveasfilename
from designs.designWord import *
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from tkinter import Tk
import docx


class WordWindow(QMainWindow, Ui_designWord):
    def __init__(self, parent=None):
        super().__init__(parent)
        super().setupUi(self)

        ### Botões da interface ###
        self.btnConverter.clicked.connect(self.convertToPdf)
        self.btnArquivo.clicked.connect(self.openFile)
        self.btnClear.clicked.connect(self.clearList)

        ### Diretório do arquivo PDF ###
        self.dir = ''

    ### Seleciona o diretório do arquivo PDF ###
    def openFile(self):
        Tk().withdraw()
        path = askopenfilename()
        
        self.dir = path
        self.listArquivos.addItem(str(path))

    def clearList(self):
        self.dir = ''
        self.listArquivos.clear()
        self.label_2.setText('')

    ### Converte arquivo PDF em arquivo Word(.docx) ###
    def convertToPdf(self):
        docWord = docx.Document()
        output_string = StringIO()

        try:
            ### Abre e lê o arquivo PDF, transformando-o em uma string ###
            with open(fr'{self.dir}', 'rb') as in_file:
                parser = PDFParser(in_file)
                doc = PDFDocument(parser)
                codec = 'utf-8'
                rsrcmgr = PDFResourceManager()
                device = TextConverter(rsrcmgr, output_string, laparams=LAParams(), codec=codec)
                interpreter = PDFPageInterpreter(rsrcmgr, device)
                
                for page in PDFPage.create_pages(doc):
                    interpreter.process_page(page)

            texto = output_string.getvalue()
            
            ### Converte o texto em caracteres válidos ###
            def valid_xml_char_ordinal(c):
                codepoint = ord(c)
        
                return (
                    0x20 <= codepoint <= 0xD7FF or
                    codepoint in (0x9, 0xA, 0xD) or
                    0xE000 <= codepoint <= 0xFFFD or
                    0x10000 <= codepoint <= 0x10FFFF
                    )

            ### String limpa ###
            cleaned_string = ''.join(c for c in texto if valid_xml_char_ordinal(c))

            ### Adiciona a string em um documento .docx ###
            docWord.add_paragraph(cleaned_string)

            ### Salva o arquivo na pasta selecionada ###
            saveDir = asksaveasfilename(defaultextension='.docx')
            docWord.save(saveDir)

            self.label_2.setText('Arquivo Word criado com sucesso!')

        except Exception as e:
            self.label_2.setText('Erro! Verifique os diretórios das imagens.')
